from PIL import Image
import pandas as pd
import os
import docx
import sys
import piexif




def reduceImgs(imgPaths, size = [512,512], overwrite = False):
    numPaths = len(imgPaths)
    reducedPaths = []
    print("Reducing Imgs")
    for i, imgPath in enumerate(imgPaths):
        imgPath = imgPath[1:-1]
        print(round(i/numPaths*100, 2))
        splitPath = imgPath.split(".")
        reducedPath = splitPath[0] + "_reduced." + splitPath[1]
        reducedPaths.append(reducedPath)
        if os.path.isfile(reducedPath) & (overwrite == False):
            continue
        else:
            date_time_original_tag = 36867
            with Image.open(imgPath) as img:
                exif_data = img._getexif()
                orig_time = exif_data[date_time_original_tag]
                exif_dict = {"Exif": {piexif.ExifIFD.DateTimeOriginal: orig_time}}
                exif_bytes = piexif.dump(exif_dict)
                img.thumbnail(size)
                img.save(reducedPath, exif = exif_bytes)

    print("Done reducing!")
    return reducedPaths

def generateReport(data):
    CURR_DIR = __file__[:-len(os.path.basename(__file__))]

    wordDocument = docx.Document()

    para = wordDocument.add_paragraph()
    pageTitle = para.add_run("Inspection Findings")
    pageTitle.font.size = docx.shared.Pt(24)
    pageTitle.bold = True
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    numImgs = len(data["imgID"])

    for i, imgID in enumerate(data["imgID"]):
        if i%2 == 0:
            table = wordDocument.add_table(rows = 4, cols=1)
            table.style = 'Table Grid'
            wordDocument.add_page_break()
        
        table.cell((i%2)*2, 0).text = imgID
        table.cell((i%2)*2+1, 0).paragraphs[0].add_run().add_picture(data["imgPaths"][i],  width=docx.shared.Inches(5))
        table.cell((i%2)*2+1, 0).paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    wordDocument.save(f"{CURR_DIR}\\test.docx")

    return

def Main():
    
    with open(r"L:\hrltec\OPS\Paiton\48241983 - Piping Inspection and Hanger Survey\Walkdowns\Script\inspectionImgs.txt", 'r') as file:
        imgNums = []
        imgPaths = []
        for line in file.readlines():
            list_pair = line.split(" ",1)
            imgNums.append(list_pair[0])
            imgPaths.append(list_pair[1].strip())
    df = pd.DataFrame()
    df["imgID"] = imgNums
    df["imgPaths"] = reduceImgs(imgPaths)
    generateReport(df)

    return

if __name__ == "__main__":
    Main()