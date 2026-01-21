import os
import sys
import pandas as pd
import openpyxl
import docx
import qrcode
import cv2
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def getData(xlDoc):

    sheets = pd.read_excel(xlDoc, sheet_name = None, header  = 0)

    data = pd.DataFrame()
    idCols = ["Line", "Side", "Hanger Designation"]

    for sheet in sheets:
        if all(x in sheets[sheet].columns for x in idCols):
            print(sheets[sheet])
            data = pd.concat([data, sheets[sheet]], join = "outer", ignore_index=True)

    data = data[data["Line"].notna() & data["Side"].notna() & data["Hanger Designation"].notna()]
    data.round(decimals=1)
    return data

def Main():

    CURR_DIR = __file__[:-len(os.path.basename(__file__))]

    hangerData = getData(sys.argv[1]).iloc[1:]
    print(hangerData)

    wordDoc = docx.Document()
    numCols = 4
    colWidth = 5
    rowsPerPage = 5
    codesPerPage = rowsPerPage*numCols
    codeSize = 1.5

    if not os.path.exists(f"{CURR_DIR}\\QRCodes"):
        os.mkdir(f"{CURR_DIR}\\QRCodes")

    
    for i, ind in enumerate(hangerData.index): 

        ID = f"{hangerData["Line"][ind]} {hangerData["Side"][ind]} {hangerData["Hanger Designation"][ind]}".replace("*","")
        print(ID)
        qrCode = qrcode.QRCode()
        qrCode.add_data(ID)
        img = qrCode.make_image()

        img.save(f"{CURR_DIR}\\QRCodes\\{ID}.png")

        if i%codesPerPage == 0:
            offset = i
            table = wordDoc.add_table(rows = rowsPerPage, cols = numCols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            wordDoc.add_page_break()

        cell = table.cell(int((i-offset)/numCols), (i-offset)%numCols)
        cell.width = docx.shared.Cm(colWidth)
        run = cell.paragraphs[0].add_run(f"{i+1}\n")
        run.add_text(ID)
        run.font.size = docx.shared.Pt(10)
        run.add_picture(f"{CURR_DIR}\\QRCodes\\{ID}.png", width=docx.shared.Inches(codeSize), height = docx.shared.Inches(codeSize))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in wordDoc.sections:
        section.top_margin = docx.shared.Cm(0.75)
        section.left_margin = docx.shared.Cm(1)
        section.bottom_margin = docx.shared.Cm(0)
        section.right_margin = docx.shared.Cm(1)

    wordDoc.save(f"{CURR_DIR}\\qrTest.docx")

    return

if __name__ == "__main__":
    Main()
    #Read()