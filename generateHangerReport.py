import docx
from PIL import Image
from PIL.ExifTags import TAGS
import openpyxl
import piexif
import pandas as pd
import sys
import os

#print(ws.cell(row = 1, column = 1).value) 

# p = wordDocument.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# wordDocument.add_heading('Heading, level 1', level=1)
# wordDocument.add_paragraph('Intense quote', style='Intense Quote')

# wordDocument.add_paragraph(
    # 'first item in unordered list', style='List Bullet'
# )
# wordDocument.add_paragraph(
    # 'first item in ordered list', style='List Number'
# )

#wordDocument.add_picture(CURR_DIR + '\\Me.JPG', width=docx.shared.Inches(1.25))

def getData(xlDoc):

    sheets = pd.read_excel(xlDoc, sheet_name = None, header  = 1)

    print(sheets)

    data = pd.DataFrame()

    for sheet in sheets:
        data = pd.concat([data, sheets[sheet]], join = "outer", ignore_index=True)

    data.round(decimals=1)

    return data

def getPictures(root_folder):

    img_paths = []
    for root, dirs, files in os.walk(root_folder):
        for file in files:
            if file.lower() in ["g.jpg", "c.jpg", "h.jpg"]:
               img_paths.append(os.path.join(root, file))
    return img_paths

def matchImgs(data, img_paths):

    data["gen_img"] = pd.Series(dtype = "str")
    data["cold_img"] = pd.Series(dtype = "str")
    data["hot_img"] = pd.Series(dtype = "str")

    for i, entry in enumerate(data["Hanger Designation"]):
        for path in img_paths:
            path_list = path.split("\\")
            ext = path_list[-1].split(".")[1]
            if data["Line"][i] in path_list and entry in path_list:
                fileName = path_list[-1].lower().split(".")[0]
                reducedPath = "\\".join(path_list[0:-1])+f"\\{fileName}_reduced.{ext}"

                if fileName == "g":
                    if pd.isna(data["gen_img"][i]):
                        data.loc[i, "gen_img"] = reducedPath
                    else:
                        print(f"Attempted to overwrite {data["gen_img"][i]} with {path}")
                if fileName == "c":
                    if pd.isna(data["cold_img"][i]):
                        data.loc[i, "cold_img"] = reducedPath
                    else:
                        print(f"Attempted to overwrite {data["cold_img"][i]} with {path}")                  
                if fileName == "h":
                    if pd.isna(data["hot_img"][i]):
                        data.loc[i, "hot_img"] = reducedPath
                    else:
                        print(f"Attempted to overwrite {data["hot_img"][i]} with {path}")

        if pd.isna(data["gen_img"][i]):
            print(f"Cannot find gen img for {str(data["Line"][i]) + " " + str(entry)}")
        if pd.isna(data["cold_img"][i]):
            print(f"Cannot find cold img for {str(data["Line"][i]) + " " + str(entry)}")
        if pd.isna(data["hot_img"][i]):
            print(f"Cannot find hot img for {str(data["Line"][i]) + " " + str(entry)}")


    return data
                
def reduceImgs(imgPaths, size = [512,512], overwrite = False):
    numPaths = len(imgPaths)
    print("Reducing Imgs")
    for i, imgPath in enumerate(imgPaths):
        print(round(i/numPaths*100, 2))
        splitPath = imgPath.split(".")
        reducedPath = splitPath[0] + "_reduced." + splitPath[1]
        if os.path.isfile(reducedPath) & (overwrite == False):
            continue
        else:
            date_time_original_tag = 36867
            with Image.open(imgPath) as img:
                exif_data = img._getexif()
                orig_time = exif_data[date_time_original_tag]
                exif_dict = {"Exif": {piexif.ExifIFD.DateTimeOriginal: orig_time}}
                exif_bytes = piexif.dump(exif_dict)
                img.thumbnail(size)
                img.save(reducedPath, exif = exif_bytes)

    print("Done reducing!")
    return

def formatDate(exifDate):
    date = exifDate.split(" ")[0].split(":")
    year = date[0]
    month = int(date[1])
    day = date[2]

    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

    return f"{day} {months[month-1]} {year}"

def GenerateTemplate(wordDocument, hangerData):

    CURR_DIR = __file__[:-len(os.path.basename(__file__))]

    abr_dict = {
        "MS" : "Main Steam",
        "HRH" : "Hot Reheat",
        "CRH 130" : "Cold Reheat 130",
        "CRH 131" : "Cold Reheat 131",
        "FW" : "Feedwater"
    }
    
    numHangers = len(hangerData["Line"])

    for i, line in enumerate(hangerData["Line"]):

        print(f"Generating document: {round(i/numHangers*100,1)}% complete")
    
        title = "Unit 8 " + abr_dict[line] + " - " + hangerData["Hanger Designation"][i]
        
        para = wordDocument.add_paragraph()
        pageTitle = para.add_run(title)
        pageTitle.font.size = docx.shared.Pt(24)
        pageTitle.bold = True
        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        #Table 1

        headings_1 = ["Type",
                    "Nameplate Units",
                    "Nameplate Cold Load",
                    "Nameplate Hot Load",
                    "Nameplate Spring Rate",
                    "Nameplate Total Travel",
                    "Site Indicator Cold",
                    "Site Indicator Hot"
        ]

        table = wordDocument.add_table(rows = 0, cols=3)
        table.style = 'Table Grid'
        
        for j, heading in enumerate(headings_1):
            if(not(heading == None)):
                row_cells = table.add_row().cells
                row_cells[0].text = heading

                #Round value if it is a number
                val = hangerData[heading][i]
                if str(val).replace(".","",1).isnumeric():
                    val = round(val,1)
                if str(val) == "nan":
                    val = "-"
                row_cells[1].text = str(val)               

        table.cell(0,2).merge(table.cell(len(table.rows)-1,2))
        
        #table 1 pic        
        table.cell(0,2).paragraphs[0].add_run().add_picture(hangerData["gen_img"][i], width=docx.shared.Inches(2), height = docx.shared.Inches(2))

        #Table 2
        headings_2 = ["Site Reading 1 Cold",
                    "Site Reading 2 Cold",
                    "Cold Notes",
        ]
        
        wordDocument.add_paragraph()
        table = wordDocument.add_table(rows = 0, cols=3)
        table.style = 'Table Grid'

        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run("Cold Survey Result").bold = True

        for j, heading in enumerate(headings_2):
                row_cells = table.add_row().cells
                row_cells[0].text = heading

                #Round value if it is a number
                val = hangerData[heading][i]
                if str(val).replace(".","",1).isnumeric():
                    val = round(val,1)
                if str(val) == "nan":
                    val = "-"
                row_cells[1].text = str(val)  

        table.cell(0,2).merge(table.cell(len(table.rows)-1,2))
        table.cell(0,0).merge(table.cell(0,1))

        #table 2 pic        
        table.cell(0,2).paragraphs[0].add_run().add_picture(hangerData["cold_img"][i], width=docx.shared.Inches(2), height = docx.shared.Inches(2))

        #Table 3
        headings_3 = ["Site Reading 1 Hot",
                    "Site Reading 2 Hot",
                    "Hot Notes",
        ]
        
        wordDocument.add_paragraph()
        table = wordDocument.add_table(rows = 0, cols=3)
        table.style = 'Table Grid'

        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run("Cold Survey Result").bold = True

        for j, heading in enumerate(headings_3):
                row_cells = table.add_row().cells
                row_cells[0].text = heading

                #Round value if it is a number
                val = hangerData[heading][i]
                if str(val).replace(".","",1).isnumeric():
                    val = round(val,1)
                if str(val) == "nan":
                    val = "-"
                row_cells[1].text = str(val)  

        table.cell(0,2).merge(table.cell(len(table.rows)-1,2))
        table.cell(0,0).merge(table.cell(0,1))

        #table 3 pic        
        table.cell(0,2).paragraphs[0].add_run().add_picture(hangerData["hot_img"][i], width=docx.shared.Inches(2), height = docx.shared.Inches(2))

        #Table 4
        headings_4 = ["Site Reading 1 Cold",
                    "Site Reading 1 Hot",
                    "Actual 1: Cold to Hot (mm)",
                    "Site Reading 2 Cold",
                    "Site Reading 2 Hot",
                    "Actual 2: Cold to Hot (mm)",
                    "General Notes"
        ]        
        data_line = 2
        
        wordDocument.add_paragraph()
        table = wordDocument.add_table(rows = 4, cols=5)
        table.style = 'Table Grid'

        #Headings on row 1
        table.cell(0,0).text = "Cold Surveys"
        table.cell(0,2).text = "Hot Surveys"
        table.cell(0,4).text = "C-H Travel (mm)"

        #Headings on row 2
        table.cell(1,0).text = "Date"
        table.cell(1,1).text = "Position"
        table.cell(1,2).text = "Date"
        table.cell(1,3).text = "Position"

        #data        

        with Image.open(hangerData["cold_img"][i]) as img:
            cold_date = img._getexif()[36867]
            table.cell(data_line,0).text = formatDate(str(cold_date))

        table.cell(data_line,1).text = str(hangerData[headings_4[0]][i])

        with Image.open(hangerData["hot_img"][i]) as img:
            hot_date = img._getexif()[36867]
            table.cell(data_line,2).text = formatDate(str(hot_date))

        table.cell(data_line,3).text = str(hangerData[headings_4[1]][i])

        c2h_travel = hangerData[headings_4[2]][i]
        if str(c2h_travel).replace(".", "",1).replace("-","",1).isnumeric():
            c2h_travel = round(c2h_travel,1)
        table.cell(data_line,4).text = str(c2h_travel)

        #Extra data for second walkdown
        cold_date_2 = "1 Feb 2025"
        hot_date_2 = "2 Mar 2025"

        table.cell(data_line+1,0).text = str(cold_date_2)
        table.cell(data_line+1,1).text = str(hangerData[headings_4[3]][i])
        table.cell(data_line+1,2).text = str(hot_date_2)
        table.cell(data_line+1,3).text = str(hangerData[headings_4[4]][i])

        c2h_travel_2 = hangerData[headings_4[5]][i]
        if str(c2h_travel_2).replace(".", "",1).replace("-","",1).isnumeric():
            c2h_travel_2 = round(c2h_travel_2,1)
        table.cell(data_line+1,4).text = str(c2h_travel_2)

        #Formatting

        table.cell(0,0).merge(table.cell(0,1))
        table.cell(0,2).merge(table.cell(0,3))
        table.cell(0,4).merge(table.cell(1,4))
        
        wordDocument.add_page_break()
        
    return wordDocument

def Main():

    CURR_DIR = __file__[:-len(os.path.basename(__file__))]

    xlFile = sys.argv[1]
    root_folder = sys.argv[2]
    print(root_folder)
    data = getData(xlFile)
    print(data)
    paths = getPictures(root_folder)
    reduceImgs(paths, overwrite=False)
    print("Got paths!")
    merged = matchImgs(data, paths)

    wordDocument = docx.Document()
    wordDocument = GenerateTemplate(wordDocument, merged)
    
    wordDocument.save(CURR_DIR + '\\demo.docx')


    print(merged)

    # CURR_DIR = __file__[:-len(os.path.basename(__file__))]
    
    # pictureRoot = CURR_DIR + "/Root"

    # try:
    #     draggedFile = sys.argv[1]
    # except:
        
    #     Exception("Drag and drop must be used")

    # xlDoc = openpyxl.load_workbook(draggedFile, data_only = True)    
    # hangerData = getData(xlDoc)    
    
    # wordDocument = docx.Document()
    # wordDocument = GenerateTemplate(wordDocument, hangerData)  
    
    # wordDocument.save(CURR_DIR + '\demo.docx')

if __name__ == "__main__":
    Main()